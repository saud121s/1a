import os
import asyncio
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import FSInputFile, ReplyKeyboardMarkup, KeyboardButton
from openai import OpenAI
from docx import Document

TOKEN = os.getenv("BOT_TOKEN")
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
bot = Bot(token=TOKEN)
dp = Dispatcher()
user_states = {}

# دالة إنشاء الملف (تنسيق أكاديمي)
def create_academic_doc(data, content, title):
    doc = Document()
    # البيانات في الأعلى
    doc.add_paragraph(f"الاسم: {data.get('name')}\nالجامعة: {data.get('uni')}\nالدكتور: {data.get('dr')}\nالرقم الجامعي: {data.get('id')}")
    doc.add_paragraph("_" * 40)
    # العنوان والمحتوى
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    file_path = f"{title.replace(' ', '_')}.docx"
    doc.save(file_path)
    return file_path

@dp.message(Command("start"))
async def start(message: types.Message):
    kb = ReplyKeyboardMarkup(keyboard=[
        [KeyboardButton(text="📸 حل اختبار")],
        [KeyboardButton(text="📝 واجب"), KeyboardButton(text="🔍 بحث")]
    ], resize_keyboard=True)
    await message.answer("أهلاً بك يا سعود. أنا مساعدك الأكاديمي، اختر الخدمة:", reply_markup=kb)

@dp.message(F.text.in_({"📸 حل اختبار", "📝 واجب", "🔍 بحث"}))
async def menu_select(message: types.Message):
    user_states[message.from_user.id] = message.text
    if message.text == "📸 حل اختبار":
        await message.answer("أرسل صورة السؤال وسأقوم بحله فوراً.")
    else:
        await message.answer(f"لـ {message.text}، أرسل بياناتك (الاسم، الجامعة، الدكتور، الرقم الجامعي) متبوعة بموضوع {message.text.replace('📝', '').replace('🔍', '')}.")

@dp.message(F.text)
async def handle_text(message: types.Message):
    state = user_states.get(message.from_user.id)
    if state in ["📝 واجب", "🔍 بحث"]:
        await message.answer("جاري التجهيز..")
        # استخراج البيانات (افتراضياً: يتم استخراجها من النص)
        prompt = f"اكتب {state} عن: {message.text}. رتبه بأسلوب أكاديمي، ابدأ بالمحتوى مباشرة."
        res = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
        
        # بيانات تجريبية (يتم استبدالها لاحقاً ببيانات المستخدم الفعلية)
        data = {"name": "سعود", "uni": "الجامعة", "dr": "دكتور المادة", "id": "00000"}
        path = create_academic_doc(data, res.choices[0].message.content, state)
        await message.answer_document(FSInputFile(path))
        os.remove(path)
        user_states.pop(message.from_user.id)

@dp.message(F.photo)
async def handle_photo(message: types.Message):
    await message.answer("جاري الحل بدقة..")
    file = await bot.get_file(message.photo[-1].file_id)
    url = f"https://api.telegram.org/file/bot{TOKEN}/{file.file_path}"
    res = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": [{"type": "text", "text": "حل الاختبار: (مقالي/اختياري). اكتب رقم السؤال والإجابة المختصرة فقط."}, {"type": "image_url", "image_url": {"url": url}}]}])
    await message.answer(res.choices[0].message.content)

async def main(): await dp.start_polling(bot)
if __name__ == "__main__": asyncio.run(main())
